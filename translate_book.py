import argparse
import os
import sys
import tempfile
import json
from copy import deepcopy
from docx import Document
from openai import OpenAI
from google.cloud import translate_v2 as translate
from google.oauth2 import service_account

# --- Load .env explicitly (always from project folder) ---
try:
    from dotenv import load_dotenv
    DOTENV_PATH = r"C:\Jessexa.com\my book\Dreamer and the Mirror\Language Translation Ebook\.env"
    load_dotenv(dotenv_path=DOTENV_PATH)
    print(f"✅ Loaded .env from {DOTENV_PATH}")
except ImportError:
    print("⚠️ python-dotenv not installed, skipping .env load")

# --- Optional Streamlit (only used for web) ---
try:
    import streamlit as st
    STREAMLIT = True
except ImportError:
    STREAMLIT = False

# --- Setup OpenAI Keys ---
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_ORG_ID = os.getenv("OPENAI_ORG_ID")
OPENAI_PROJECT_ID = os.getenv("OPENAI_PROJECT_ID")

print("---- ENV CHECK ----")
print("OPENAI_API_KEY:", (OPENAI_API_KEY[:15] + "...") if OPENAI_API_KEY else None)
print("OPENAI_ORG_ID:", OPENAI_ORG_ID)
print("OPENAI_PROJECT_ID:", OPENAI_PROJECT_ID)
print("GOOGLE_APPLICATION_CREDENTIALS:", os.getenv("GOOGLE_APPLICATION_CREDENTIALS"))
print("-------------------")

client = None
if OPENAI_API_KEY:
    try:
        if OPENAI_API_KEY.startswith("sk-proj"):
            client = OpenAI(
                api_key=OPENAI_API_KEY,
                organization=OPENAI_ORG_ID,
                project=OPENAI_PROJECT_ID
            )
            print("✅ OpenAI client initialized with project key.")
        else:
            client = OpenAI(api_key=OPENAI_API_KEY)
            print("✅ OpenAI client initialized with classic key.")
    except Exception as e:
        print(f"❌ Failed to initialize OpenAI client: {e}")
else:
    print("❌ OpenAI API key not found (set OPENAI_API_KEY in .env or Streamlit Secrets).")

# --- Google Cloud Credentials ---
google_creds = None
try:
    if os.getenv("GOOGLE_CLOUD_KEY"):
        gcloud_key = json.loads(os.getenv("GOOGLE_CLOUD_KEY"))
        google_creds = service_account.Credentials.from_service_account_info(gcloud_key)
        print("✅ Google Cloud credentials loaded from GOOGLE_CLOUD_KEY.")
    elif os.getenv("GOOGLE_APPLICATION_CREDENTIALS"):
        google_creds = service_account.Credentials.from_service_account_file(
            os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
        )
        print(f"✅ Google Cloud credentials loaded from file: {os.getenv('GOOGLE_APPLICATION_CREDENTIALS')}")
    else:
        print("❌ Google Cloud credentials not found (set GOOGLE_APPLICATION_CREDENTIALS or GOOGLE_CLOUD_KEY).")
except Exception as e:
    print(f"❌ Error loading Google Cloud credentials: {e}")

# --- Translation Functions ---
def translate_openai(text, target_lang, mode="both"):
    if not text.strip() or not client:
        return text, text
    prompt = f"""
    Translate the following passage from English into {target_lang}.

    ### Literal
    (close to the words)

    ### Poetic
    (soulful, rhythmic, poetic, but faithful)

    Passage:
    {text}
    """
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
    )
    content = resp.choices[0].message.content

    literal, poetic = "", ""
    if "### Literal" in content and "### Poetic" in content:
        parts = content.split("### Poetic")
        literal = parts[0].replace("### Literal", "").strip()
        poetic = parts[1].strip()
    else:
        literal = content.strip()
        poetic = content.strip()

    if mode == "literal":
        return literal, ""
    elif mode == "poetic":
        return "", poetic
    return literal, poetic

def translate_google(text, target_lang):
    if not text.strip():
        return text
    if google_creds:
        translate_client = translate.Client(credentials=google_creds)
    else:
        translate_client = translate.Client()
    result = translate_client.translate(text, target_language=target_lang)
    return result["translatedText"]

# --- Format-preserving translation ---
def translate_book(input_file, provider, target_lang, mode):
    original = Document(input_file)

    doc_literal = deepcopy(original)
    doc_poetic = deepcopy(original)

    for i, para in enumerate(original.paragraphs):
        if para.text.strip():
            if provider == "openai":
                literal, poetic = translate_openai(para.text, target_lang, mode)
                if literal:
                    doc_literal.paragraphs[i].text = literal
                if poetic:
                    doc_poetic.paragraphs[i].text = poetic
            elif provider == "google":
                new_text = translate_google(para.text, target_lang)
                doc_literal.paragraphs[i].text = new_text
        else:
            doc_literal.paragraphs[i].text = ""
            doc_poetic.paragraphs[i].text = ""

    return doc_literal, doc_poetic

# --- CLI Mode ---
def cli_main():
    parser = argparse.ArgumentParser(description="Translate a whole book (.docx) with formatting")
    parser.add_argument("input", help="Input .docx file")
    parser.add_argument("output", help="Output .docx file")
    parser.add_argument("language", help="Target language code (e.g. fr, vi, es)")
    parser.add_argument("--provider", choices=["openai", "google"], default="openai")
    parser.add_argument("--mode", choices=["literal", "poetic", "both"], default="both",
                        help="For OpenAI only")
    args = parser.parse_args()

    doc_literal, doc_poetic = translate_book(args.input, args.provider, args.language, args.mode)

    if args.provider == "openai":
        if args.mode in ["literal", "both"]:
            doc_literal.save(args.output.replace(".docx", "_literal.docx"))
        if args.mode in ["poetic", "both"]:
            doc_poetic.save(args.output.replace(".docx", "_poetic.docx"))
    else:
        doc_literal.save(args.output)

    print("✅ Translation complete! Formatting preserved.")

# --- Web Mode ---
def web_main():
    st.title("📖 Language Translation eBook")
    uploaded_file = st.file_uploader("Upload your .docx file", type="docx")

    provider = st.selectbox("Translation provider", ["openai", "google"])
    target_lang = st.text_input("Target language code (e.g. fr, vi, es)", "vi")
    mode = st.selectbox("Mode (for OpenAI only)", ["both", "literal", "poetic"])

    if uploaded_file and st.button("Translate"):
        try:
            st.info("⚡ Starting translation... please wait")

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(uploaded_file.read())
                tmp_path = tmp.name

            doc_literal, doc_poetic = translate_book(tmp_path, provider, target_lang, mode)

            st.success("✅ Translation finished! Formatting preserved.")

            if provider == "openai":
                if mode in ["literal", "both"]:
                    lit_path = "translated_literal.docx"
                    doc_literal.save(lit_path)
                    with open(lit_path, "rb") as f:
                        st.download_button("⬇️ Download Literal", f, file_name=lit_path)

                if mode in ["poetic", "both"]:
                    poe_path = "translated_poetic.docx"
                    doc_poetic.save(poe_path)
                    with open(poe_path, "rb") as f:
                        st.download_button("⬇️ Download Poetic", f, file_name=poe_path)
            else:
                out_path = "translated.docx"
                doc_literal.save(out_path)
                with open(out_path, "rb") as f:
                    st.download_button("⬇️ Download Translation", f, file_name=out_path)
        except Exception as e:
            st.error(f"❌ Something went wrong: {e}")

# --- Entry Point ---
if __name__ == "__main__":
    if len(sys.argv) > 1 and not sys.argv[1].startswith("-"):
        cli_main()
    elif STREAMLIT:
        web_main()
    else:
        print("ℹ️ Run with arguments for CLI, or `streamlit run translate_book.py` for web UI")
